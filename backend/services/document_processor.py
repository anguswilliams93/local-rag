import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import settings
from storage.file_storage import file_storage
from storage.vector_store import vector_store

logger = logging.getLogger("local-rag.processor")


class DocumentProcessor:
    """Process documents of various formats into chunks for RAG."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv", ".xlsx", ".xls"}

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported."""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    def get_file_type(self, filename: str) -> str:
        """Get normalized file type from filename."""
        ext = Path(filename).suffix.lower()
        type_map = {
            ".pdf": "pdf",
            ".txt": "text",
            ".md": "markdown",
            ".docx": "docx",
            ".csv": "csv",
            ".xlsx": "excel",
            ".xls": "excel"
        }
        return type_map.get(ext, "unknown")

    def extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text content from a file."""
        logger.info(f"📄 Extracting text: type={file_type}, path={file_path.name}")

        if file_type == "pdf":
            text = self._extract_pdf(file_path)
        elif file_type in ("text", "markdown"):
            text = self._extract_text_file(file_path)
        elif file_type == "docx":
            text = self._extract_docx(file_path)
        elif file_type in ("csv", "excel"):
            text = self._extract_tabular(file_path, file_type)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        logger.info(f"✅ Extracted {len(text)} characters")
        return text

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(str(file_path))
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {page_num}]\n{page_text}")

        return "\n\n".join(text_parts)

    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from TXT or Markdown file."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue

        # Fallback: read as bytes and decode with errors='replace'
        return file_path.read_bytes().decode("utf-8", errors="replace")

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(str(file_path))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                paragraphs.append("\n".join(table_text))

        return "\n\n".join(paragraphs)

    def _extract_tabular(self, file_path: Path, file_type: str) -> str:
        """Extract text from CSV or Excel file with table-aware formatting."""
        try:
            if file_type == "csv":
                df = pd.read_csv(file_path)
            else:  # excel
                df = pd.read_excel(file_path)

            # For small tables, convert entire table to text
            if len(df) <= 50:
                return self._dataframe_to_text(df)

            # For larger tables, chunk by row groups while preserving headers
            return self._dataframe_to_text(df)
        except Exception as e:
            raise ValueError(f"Error processing tabular file: {str(e)}")

    def _dataframe_to_text(self, df: pd.DataFrame) -> str:
        """Convert a pandas DataFrame to a readable text format."""
        return df.to_string(index=False)

    def process_document(self, agent_id: str, file_path: Path, original_filename: str) -> int:
        """Process a document into chunks and add to vector store."""
        logger.info(f"⚙️  Processing: {original_filename}")

        file_type = self.get_file_type(original_filename)
        text = self.extract_text(file_path, file_type)

        logger.info(f"✂️  Chunking text (size={settings.chunk_size}, overlap={settings.chunk_overlap})...")
        chunks = self.text_splitter.split_text(text)
        logger.info(f"✅ Created {len(chunks)} chunks")

        metadatas = [
            {
                "source": original_filename,
                "agent_id": agent_id,
                "chunk_index": i,
                "timestamp": datetime.utcnow().isoformat()
            }
            for i in range(len(chunks))
        ]

        logger.info(f"💾 Adding chunks to vector store...")
        vector_store.add_documents(agent_id, chunks, metadatas)
        logger.info(f"✅ Document processed successfully")
        return len(chunks)


# Singleton instance
document_processor = DocumentProcessor()
